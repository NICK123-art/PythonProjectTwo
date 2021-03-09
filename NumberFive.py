import docx

wdoc = docx.Document()
wdoc.add_heading('明智可定')
wdoc.add_heading('明智可定',level=2)
wdoc.add_heading('明智可定',level=3)
print(wdoc.paragraphs[0].text)
print(wdoc.paragraphs[1].text)
print(wdoc.paragraphs[2].text)
wdoc.save('out17_4.docx')