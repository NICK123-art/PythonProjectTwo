import docx
wdoc = docx.Document('data17_1.docx')
print("段落数是 = ", len(wdoc.paragraphs))
for i in range(0, len(wdoc.paragraphs)):
    print("paragraph %d = "%i, wdoc.paragraphs[i].text)

print("Run数量 = ",len(wdoc.paragraphs[1].runs))