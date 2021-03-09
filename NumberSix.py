import docx
wdoc = docx.Document()
ptr = wdoc.add_paragraph('我是第一段落')
run1 = ptr.add_run('我是斜体')
run1.bold = True
run2 = ptr.add_run('我是粗体')
run2.italic = True
wdoc.save('out14_15.docx')